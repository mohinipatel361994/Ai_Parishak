import streamlit as st
from zipfile import ZipFile
from pdfminer.pdfinterp import PDFResourceManager, PDFPageInterpreter
from pdfminer.converter import TextConverter
from pdfminer.layout import LAParams
from pdfminer.pdfpage import PDFPage
from docx import Document 
import pandas as pd

from io import StringIO
import base64
#------- OCR ------------
import pdf2image
import pytesseract
from pytesseract import Output, TesseractError

#@st.cache_data
# def images_to_txt(path, language):
#     images = pdf2image.convert_from_bytes(path,poppler_path= r'C:\Users\YH185MX\Downloads\poppler-23.11.0\Library\bin')
#     all_text = []
#     for i in images:
#         pil_im = i
#         text = pytesseract.image_to_string(pil_im, lang=language)
#         # ocr_dict = pytesseract.image_to_data(pil_im, lang='eng', output_type=Output.DICT)
#         # ocr_dict now holds all the OCR info including text and location on the image
#         # text = " ".join(ocr_dict['text'])
#         # text = re.sub('[ ]{2,}', '\n', text)
#         all_text.append(text)
#     text='/n'.join(all_text)
#     print(type(text))
#     return text, len(all_text)

def images_to_txt(path, language):
    images = pdf2image.convert_from_bytes(path,poppler_path= r'C:\poppler-24.02.0\Library\bin')
    all_text = []
    for v,i in enumerate(images):
        pil_im = i
        text = f"\nPage_Number {v} started\n"
        text += pytesseract.image_to_string(pil_im, lang='eng+pan')
        text+= f"\nPage_Number {v} ended\n"
        # ocr_dict = pytesseract.image_to_data(pil_im, lang='eng', output_type=Output.DICT)
        # ocr_dict now holds all the OCR info including text and location on the image
        # text = " ".join(ocr_dict['text'])
        # text = re.sub('[ ]{2,}', '\n', text)
        all_text.append(text)
    text='/n'.join(all_text)
    print(type(text))
    return text, len(all_text)

@st.cache_data
def convert_pdf_to_txt_pages(path):
    texts = []
    rsrcmgr = PDFResourceManager()
    retstr = StringIO()
    laparams = LAParams()
    device = TextConverter(rsrcmgr, retstr, laparams=laparams)
    # fp = open(path, 'rb')
    interpreter = PDFPageInterpreter(rsrcmgr, device)
    size = 0
    c = 0
    file_pages = PDFPage.get_pages(path)
    nbPages = len(list(file_pages))
    for page in PDFPage.get_pages(path):
      interpreter.process_page(page)
      t = retstr.getvalue()
      if c == 0:
        texts.append(t)
      else:
        texts.append(t[size:])
      c = c+1
      size = len(t)
    # text = retstr.getvalue()

    # fp.close()
    device.close()
    retstr.close()
    return texts, nbPages

@st.cache_data
def convert_pdf_to_txt_file(path):
    texts = []
    rsrcmgr = PDFResourceManager()
    retstr = StringIO()
    laparams = LAParams()
    device = TextConverter(rsrcmgr, retstr, laparams=laparams)
    # fp = open(path, 'rb')
    interpreter = PDFPageInterpreter(rsrcmgr, device)
    
    file_pages = PDFPage.get_pages(path)
    nbPages = len(list(file_pages))
    for page in PDFPage.get_pages(path):
      interpreter.process_page(page)
      t = retstr.getvalue()
    # text = retstr.getvalue()

    # fp.close()
    device.close()
    retstr.close()
    return t, nbPages

@st.cache_data
def save_pages(pages):
  
  files = []
  for page in range(len(pages)):
    filename = "page_"+str(page)+".txt"
    with open("./file_pages/"+filename, 'w', encoding="utf-8") as file:
      file.write(pages[page])
      files.append(file.name)
  
  # create zipfile object
  zipPath = './file_pages/pdf_to_txt.zip'
  zipObj = ZipFile(zipPath, 'w')
  for f in files:
    zipObj.write(f)
  zipObj.close()

  return zipPath

def displayPDF(file):
  # Opening file from file path
  # with open(file, "rb") as f:
  base64_pdf = base64.b64encode(file).decode('utf-8')

  # Embedding PDF in HTML
  pdf_display = F'<iframe src="data:application/pdf;base64,{base64_pdf}" width="700" height="1000" type="application/pdf"></iframe>'
  # Displaying File
  st.markdown(pdf_display, unsafe_allow_html=True)


def read_word_table(file_path):
  document = Document(file_path)
  tables = []
  for table in document.tables:
      data = []
      for row in table.rows:
          row_data = []
          for cell in row.cells:
              row_data.append(cell.text)
          data.append(row_data)
      df = pd.DataFrame(data[1:], columns=data[0])  
      tables.append(df)
  return tables
  

# def read_word_table(file_path):
#     document = Document(file_path)
#     tables = []
#     for table in document.tables:
#         data = []
        
#         # num_cols = len(table.columns)
#         # columns = [f"Column_{i+1}" for i in range(num_cols)]
#         columns = ['S.No','English','Punjabi','col4','col5']
        
#         for row in table.rows:
#             row_data = [cell.text for cell in row.cells]
#             data.append(row_data)
#         df = pd.DataFrame(data, columns=columns)  
#         tables.append(df)
#     return tables
  
def create_word_doc(text):
  doc = Document()
  doc.add_paragraph(text)
  return doc
  
def get_text(filename):
  doc = Document(filename)
  full_text = []
  for para in doc.paragraphs:
      full_text.append(para.text)
  return '\n'.join(full_text)
