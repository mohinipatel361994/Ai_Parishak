import json
import pdfplumber
from langchain_community.llms import OpenAI
import chat
import openai
from langchain_community.llms import OpenAI
import streamlit as st
import os
import pandas as pd
from rag import *
from PIL import Image
from chat import load_chain
import numpy as np
import logging
from prompts import initialise_prompt
import pytesseract
from pdf2image import convert_from_bytes
from PyPDF2 import PdfReader
from langchain_openai import ChatOpenAI
from langchain.chains import ConversationChain
from langchain.memory import ConversationBufferMemory
import docx
from prompts import ai_prompt,ai_topic_prompt,latex_prompt,student_prompt,mcq_test_prompt,learn_outcome_prompt,ai_topic_prompt1
from functions import read_word_table,create_word_doc,get_text
import subprocess
import re
import constants
from dotenv import load_dotenv
from docxlatex import Document
import json
import base64
import fixed_function
from io import BytesIO
import streamlit as st
from PyPDF2 import PdfReader
from langchain_community.embeddings import OpenAIEmbeddings
from langchain.text_splitter import CharacterTextSplitter
from langchain.chains.question_answering import load_qa_chain
from langchain.llms import OpenAI
from langchain_community.vectorstores import FAISS
from docx import Document
from translate import Translator
import os
import time
from io import StringIO
#from docxlatex import Document
# from PIL import Image
#c = OpenAI()

st.set_page_config(layout='wide')

hide_streamlit_style = """
            <style>
            
            footer {visibility: hidden;}
            h1 {
               color: green;
               font-size: 35px;    
               }
            
            </style>
            """
st.markdown(hide_streamlit_style, unsafe_allow_html=True)
#add_selectbox = st.sidebar.selectbox(
#    "Store the chat hsitory",
#    ("Teachers", "Students", "Administration")
#)

load_dotenv()

images = ['6MarkQ']


#os.environ["OPENAI_API_TYPE"] = ""
#os.environ["OPENAI_API_VERSION"] = ""openai_api_key = os.getenv("OPENAI_API_KEY")

openai_api_key2 = st.secrets["secret_section"]["OPENAI_API_KEY"]

#openai_api_key = os.getenv('OPENAI_API_KEY')
TEMP_MD_FILE = r"question.pdf"
TEMP_PDF_FILE = "__temp.pdf"

def get_base64(bin_file):
    with open(bin_file, 'rb') as f:
        data = f.read()
    return  base64.b64encode(data).decode()



def download_doc(doc):
    doc_buffer = BytesIO()
    doc.save(doc_buffer)
    doc_buffer.seek(0)
    return doc_buffer


with open('style.css') as f:
 st.markdown(f'<style>{f.read()}</style>', unsafe_allow_html=True) 


def mcq_to_markdown(mcq_dict):
    markdown = ""
    for key, value in mcq_dict.items():
        if key.isdigit():
            markdown += f"Question {key}: {value}\n"
        elif key.startswith('image'):
            markdown += f"![Image {key[-1]}]({value})\n\n"
    return markdown


def display_mcq(mcq_dict):
    for key, value in mcq_dict.items():
        st.markdown(f"""**Question {key}:** {value['question']}""")
        c1,c2 = st.columns([3,1])
        with c1:
            st.write("Options:")
            for option_key, option_value in value['options'].items():
                st.write(f"{option_key}. {option_value}")
        with c2:
            st.image(f"{value['image']}", use_column_width=True)

def is_word_in_text(word, text):
    """
    Check if a word is within a given text.

    Args:
    - word (str): The word to check for.
    - text (str): The text to search within.

    Returns:
    - bool: True if the word is found in the text, False otherwise.
    """
    # Split the text into individual words
    words_in_text = text.split()

    # Check if the word is in the list of words from the text
    if word in words_in_text:
        return True
    else:
        return False
    
def text_to_latex(question):
    # Convert question to LaTeX format
    latex_question = r"\text{" + question + r"}"
    return f"{latex_question}"
    
def add_dollar_signs(text):
    # Regular expression pattern to find equations
    equation_pattern = r"([+-]?\s*\d*\s*[xX]\^\d+)"

    # Find equations in the text using regular expression
    equations = re.findall(equation_pattern, text)

    # Replace each equation with the same equation surrounded by dollar signs
    for equation in equations:
        text = text.replace(equation, f"${equation}$")

    return text
    
def generate_quiz(text):

    test_source_text = text
    quiz_rag = RAGTask(task_prompt_builder=revision_quiz_json_builder)
#     summary_rag = RAGTask(task_prompt_builder=plaintext_summary_builder)
#     glossary_rag = RAGTask(task_prompt_builder=get_glossary_builder)

    outputs = []
    # for rag_task in [quiz_rag, summary_rag, glossary_rag]:
    #     output = rag_task.get_output(source_text=test_source_text)
    #     outputs.append(output)
    for rag_task in [quiz_rag]:
        output = rag_task.get_output(source_text=test_source_text)
        outputs.append(output)
    return outputs

def decrement_question_num():
    if st.session_state['curr_question'] > 0:
        st.session_state['curr_question'] -= 1
        #ClearAll()

def increment_question_num():
    print('Incrementing question', st.session_state['curr_question'])
    if st.session_state['curr_question'] < st.session_state['quiz_length'] - 1:
        st.session_state['curr_question'] += 1
        #ClearAll()
    


def display_q(df):
    for index,row in df.iterrows():
        st.markdown("##### **Question "+str(index+1)+"**")
        if row['Image Path']=='no_image':
            # st.session_state.format_chain = ConversationChain( llm=AzureChatOpenAI(
            # deployment_name="gpt-4turbo",
            # temperature=0
            # ))
            # formatted_question = chat.format_response(row['Question'])
            # st.latex(formatted_question)
            st.info(row['Question'])
        if row['Image Path'] !='no_image':
            # st.session_state.format_chain = ConversationChain( llm=AzureChatOpenAI(
            # deployment_name="gpt-4turbo",
            # temperature=0
            # ))
            # formatted_question = chat.format_response(row['Question'])
            # st.latex(formatted_question)
            st.info(row['Question'])
            st.image(row['Image Path'],width=180)
        st.divider()

pd.set_option('display.max_colwidth', None)
st.session_state.ques = pd.read_csv('update_question.csv')



with st.container():
    col1, col2, col3 = st.columns([0.2, 0.6, 0.2], gap="small")
    with col1:
        logo_image = Image.open('assests/madhya-pradesh-logo.png')
        resized_logo = logo_image.resize((150, 150))
        st.image(resized_logo)
    with col2:
        st.markdown("# GyanKosh")
        st.markdown("###### AI Based Question Generation Assistance")
    with col3:
        l = Image.open('assests/28072020125926mpsedclogo.png')
        re = l.resize((165, 127))  # Corrected the resize method call
        st.image(re)
         
        

def on_text_area_change():
    st.session_state.page_text = st.session_state.my_text_area

def read_pdf_page(file, page_number):
    pdfReader = PdfReader(file)
    page = pdfReader.pages[page_number]
    return page.extract_text()




def markdown_to_pdf(markdown: str, output_file: str):
    """
    Convert Markdown to PDF
    :param markdown: Markdown string
    :param output_file: Output file
    """
    with open(TEMP_MD_FILE, "w",encoding='utf-8') as f:
        f.write(markdown)


def correct_bhashni_translations(text,lowercase_dict):
    corrected_text = []
    words = text.split()  # Tokenize the input text into words
    for word in words:
        # Check if the word needs correction
        if word in lowercase_dict:
            corrected_text.append(lowercase_dict[word])
        else:
            corrected_text.append(word)  # If no correction needed, keep the word unchanged
    return ' '.join(corrected_text)

def split_text_into_chunks(text, max_length):
    words = text.split()
    chunks = []
    current_chunk = []
    current_length = 0
    for word in words:
        if current_length + len(word) + 1 <= max_length:
            current_chunk.append(word)
            current_length += len(word) + 1
        else:
            chunks.append(" ".join(current_chunk))
            current_chunk = [word]
            current_length = len(word) + 1
    if current_chunk:
        chunks.append(" ".join(current_chunk))
    return chunks
from PyPDF2 import PdfReader
def list_files(folder_path):
    return os.listdir(folder_path)

# Function to remove file extension
def remove_extension(filename):
    return os.path.splitext(filename)[0]

# Function to convert PDF to text
def pdf_to_text(file_path):
    text = ""
    with open(file_path, 'rb') as file:
        pdf_reader = PdfReader(file)
        for page in pdf_reader.pages:
            text += page.extract_text()
    return text
def list_files(folder_path):
    return os.listdir(folder_path)

def remove_extension(filename):
    return os.path.splitext(filename)[0]

# Initialize session state variables
if 'selected_option' not in st.session_state:
    st.session_state.selected_option = None

if 'quesai' not in st.session_state:
    st.session_state.quesai = None

if 'selected_file' not in st.session_state:
    st.session_state.selected_file = "Select document"
   
st.sidebar.header("Select Module")
st.session_state.teach = st.sidebar.selectbox(
    "",
    ('Teachers', 'Students', 'Administration'),
    key='airadio1'
)
if st.session_state.teach == 'Teachers':
    st.session_state.quesai = st.title("Generate Question and Answer")

    if st.session_state.quesai:
        st.session_state.selected_option = st.radio(
            "Select Options",
            ("Topic Based Questions", "Text Analyzer","Pre Uploaded" , "Terminologies and Keyterms", "Learning Outcomes"),
            horizontal=True,
            index=0,
            key='option'
        )

        choose = st.session_state.selected_option

        col_11, col_22 = st.columns([2, 1])
        def pdf_to_text(pdf_path):
            text = ""
            try:
                with pdfplumber.open(pdf_path) as pdf:
                    for page in pdf.pages:
                        text += page.extract_text() or ""  # Use empty string if text extraction fails
            except Exception as e:
                st.error(f"Error reading PDF: {e}")
            return text

        with col_11:
            if choose == "Pre Uploaded":
                subjects_folder = "./preuploaded"

                # List subjects
                subjects_list = [d for d in os.listdir(subjects_folder) if os.path.isdir(os.path.join(subjects_folder, d))]
                subjects_list.sort()
                subjects_list.insert(0, "Select subject")

                selected_subject = st.selectbox("Select a subject", subjects_list, index=0, key='subject_selector')

                if selected_subject != "Select subject":
                    # List chapters for the selected subject
                    folder_path = os.path.join(subjects_folder, selected_subject)
                    files_list = list_files(folder_path)
                    files_list = [remove_extension(filename) for filename in files_list]
                    files_list.sort()
                    files_list.insert(0, "All Chapters")

                    selected_file = st.selectbox("Select a chapter (optional)", files_list, index=0, key='chapter_selector')

                    # Initialize text variable
                    st.session_state.text = ""

                    if selected_file == "All Chapters":
                        for file in files_list[1:]:  # Skip "All Chapters" option
                            pdf_file_path = os.path.join(folder_path, file + '.pdf')
                            st.session_state.text += pdf_to_text(pdf_file_path) + "\n"
                    elif selected_file != "Select chapter":
                        st.session_state.filename = []
                        pdf_file_path = os.path.join(folder_path, selected_file + '.pdf')
                        st.session_state.filename.append(selected_file)
                        st.session_state.text = pdf_to_text(pdf_file_path)

                    if st.session_state.text:
                        # Input settings
                        col1, col2 = st.columns(2)
                        with col1:
                            st.session_state.complexity = st.selectbox('Complexity Mode Required?*', ['Easy', 'Difficult'], index=0, key="mode")
                            st.session_state.no_of_questions = st.number_input('No. of Questions to generate*', key="ai_questions", step=1, max_value=30)
                            st.session_state.mode_of_questions = st.selectbox('Choose Answer Required/Not*', ['Only Questions', 'Questions with Answers'], index=0, key="quesansw")
                        with col2:
                            st.session_state.type_of_questions = st.selectbox('Choose Question Type*', ['Short Questions', 'Long Questions', 'MCQ', 'Fill in the Blanks', 'True and False'], index=0)
                            st.session_state.language = st.selectbox('Choose Response Language Mode*', ['English', 'English and Hindi'], index=0, key="lang")

                        # Button to trigger processing
                        if st.button("Submit"):
                            if st.session_state.text and st.session_state.mode_of_questions != 'Select Option':
                                st.session_state.llm = ConversationChain(llm=ChatOpenAI(model="gpt-4o", temperature=0.7, api_key=openai_api_key2))
                                chapter_info = f"Chapter: {selected_file}" if selected_file != "All Chapters" else "All Chapters"
                                formatted_output = st.session_state.llm.predict(input=ai_topic_prompt1.format(
                                    chapter_info,
                                    st.session_state.no_of_questions,
                                    st.session_state.text,
                                    st.session_state.language,
                                    st.session_state.mode_of_questions,
                                    st.session_state.type_of_questions,
                                    st.session_state.complexity,
                                    st.session_state.no_of_questions
                                ))

                                st.info(formatted_output)
                                markdown_to_pdf(formatted_output, 'question.pdf')
                                word_doc = create_word_doc(formatted_output)
                                doc_buffer = download_doc(word_doc)

                                st.download_button(
                                    label="Download Word Document",
                                    data=doc_buffer,
                                    file_name="generated_document.docx",
                                    mime="application/octet-stream",
                                    key='worddownload'
                                )

            if choose == "Terminologies and Keyterms":
                subjects_folder = "./preuploaded"
                # List subjects
                subjects_list = [d for d in os.listdir(subjects_folder) if os.path.isdir(os.path.join(subjects_folder, d))]
                subjects_list.sort()
                subjects_list.insert(0, "Select subject")

                selected_subject = st.selectbox("Select a subject", subjects_list, index=0, key='subject_selector')

                if selected_subject != "Select subject":     
                    # List chapters for the selected subject
                    folder_path = os.path.join(subjects_folder, selected_subject)                               
                    files_list = list_files(folder_path)
                    files_list = [remove_extension(filename) for filename in files_list]
                    files_list.insert(0, "Select document")
                    selected_file = st.selectbox("Select a file", files_list, index=0, key='terminologies_selected_file')

                    def mcq_response(result: str, persist_directory: str = constants.PERSIST_DIR) -> str:
                        formatted_response = st.session_state.mcq_chain.predict(input=mcq_test_prompt.format(result))
                        st.write(formatted_response)
                        return formatted_response

                    if selected_file != "Select document":
                        st.session_state.filename = []
                        with open(os.path.join(folder_path, selected_file + '.pdf'), 'rb') as file:
                            reader = PdfReader(file)
                            text = StringIO()
                            for page in reader.pages:
                                text.write(page.extract_text())
                            text = text.getvalue()
                        st.session_state.filename.append(selected_file)
                        st.session_state.mcq_chain = ConversationChain(llm=ChatOpenAI(model="gpt-4o-mini", temperature=0.7,api_key=openai_api_key2))
                        outputs = mcq_response(text)
                        markdown_to_pdf(outputs, 'question.pdf')
                        word_doc = create_word_doc(outputs)
                        doc_buffer = download_doc(word_doc)
                        st.download_button(label="Download Word Document", data=doc_buffer, file_name="generated_document.docx", mime="application/octet-stream", key='worddownload3')

            if choose == "Learning Outcomes":
                subjects_folder = "./preuploaded"
                # List subjects
                subjects_list = [d for d in os.listdir(subjects_folder) if os.path.isdir(os.path.join(subjects_folder, d))]
                subjects_list.sort()
                subjects_list.insert(0, "Select subject")

                selected_subject = st.selectbox("Select a subject", subjects_list, index=0, key='subject_selector')

                if selected_subject != "Select subject":
                    # List chapters for the selected subject
                    folder_path = os.path.join(subjects_folder, selected_subject)                
                    files_list = list_files(folder_path)
                    files_list = [remove_extension(filename) for filename in files_list]
                    files_list.insert(0, "Select document")
                    selected_file = st.selectbox("Select a file", files_list, index=0, key='learning_outcomes_selected_file')

                    def learn_outcome_term(result: str, persist_directory: str = constants.PERSIST_DIR) -> str:
                        formatted_response = st.session_state.learn_outcome_chain.predict(input=learn_outcome_prompt.format(result))
                        st.write(formatted_response)
                        return formatted_response

                    if selected_file != "Select document":
                        st.session_state.filename = []
                        with open(os.path.join(folder_path, selected_file + '.pdf'), 'rb') as file:
                            reader = PdfReader(file)
                            text = StringIO()
                            for page in reader.pages:
                                text.write(page.extract_text())
                            text = text.getvalue()
                        st.session_state.filename.append(selected_file)
                        st.session_state.learn_outcome_chain = ConversationChain(llm=ChatOpenAI(model="gpt-4o-mini", temperature=0.7,api_key=openai_api_key2))
                        outputs = learn_outcome_term(text)
                        markdown_to_pdf(outputs, 'question.pdf')
                        word_doc = create_word_doc(outputs)
                        doc_buffer = download_doc(word_doc)
                        st.download_button(label="Download Word Document", data=doc_buffer, file_name="generated_document.docx", mime="application/octet-stream", key='worddownload3')

            if choose=="Text Analyzer":
                txt = st.text_area(
                "Text to Generate Questions"
                )
                if len(txt)>0:
                    #st.write(txt)
                    with open('dictionary.json','r') as f:
                            existing_dictionary = json.load(f)

                    lowercase_dict = {key.lower(): value for key, value in existing_dictionary.items()}
                    st.session_state.text=correct_bhashni_translations(txt,lowercase_dict)
                    #st.write(st.session_state.text)

                    col1, col2 = st.columns(2)
                    with col1:
                        #st.session_state.complexity =  st.selectbox('Complexity Mode Required?*', ['No', 'Yes'],index=0,key="mode")
                        st.session_state.no_of_questions = st.number_input('No. of  Questions to generate*',key="ai_questions_no_k",step=1,max_value=30)
                        st.session_state.mode_of_questions = st.selectbox('Choose Answer Required/Not*', ['Select Option','Only Questions', 'Questions with Answers'],index=0,key="quesansw")
                    with col2:
                        st.session_state.type_of_questions =  st.selectbox('Choose Question Type*', ['Select Option','Short Questions','MCQ','Fill in the Blanks','True and False'],index=0,key="ai_questions_no_s")
                    if st.button("Submit"):
                       if st.session_state.text and st.session_state.no_of_questions>0:
                            st.session_state.llm = ConversationChain( llm=ChatOpenAI(
                            model="gpt-4o-mini",
                            temperature=0.7,
                            api_key=openai_api_key2
                            )) 
                            formatted_output = st.session_state.llm.predict(input = ai_prompt.format(st.session_state.no_of_questions,
                                                                            st.session_state.mode_of_questions,
                                                                            st.session_state.type_of_questions,
                                                                            st.session_state.text))
                            st.info(formatted_output)
                            markdown_to_pdf(formatted_output,'question.pdf')
                            word_doc = create_word_doc(formatted_output)
                            doc_buffer = download_doc(word_doc)
                            st.session_state.doc_buffer = doc_buffer
                            if 'doc_buffer' in st.session_state:
                                st.download_button(label="Download Word Document", 
                                            data=doc_buffer, 
                                            file_name="generated_document.docx", 
                                            mime="application/octet-stream",
                                            key='worddownload2')

            else:
                    st.write("")

            if choose=="Image Analyzer":
                #openai.api_type = ""
                #openai.api_version = ""
                #openai.api_base = ""  
                #openai.api_key = 
                #OPENAI_API_KEY2 =
                #openai_api_key2 = st.secrets["secret_section"]["OPENAI_API_KEY"]

                def load_image(img):
                    im = Image.open(img)
                    im = im.resize((400, 300))
                    image = np.array(im)
                    return image

                def to_base64(uploaded_file):
                    if uploaded_file is not None:
                       file_buffer = uploaded_file.read()
                       b64 = base64.b64encode(file_buffer).decode()
                       return f"data:image/png;base64,{b64}"
                    return None

                def generate_questions(image_base64):
                    response = openai.ChatCompletion.create(
                      model="gpt-3.5-turbo",
                      messages=[
                          {
                            "role": "user",
                             "content": [
                                 {"type": "text", "text": st.session_state.text_prompt},
                                 {
                                       "type": "image_url",
                                       "image_url": {
                                        "url": image_base64,
                                        },
                                 },
                              ],
                           }
                    ],
                    max_tokens=2000,
                )
                    return response.choices[0].message.content

                st.write("Upload Image to Generate Questions")
                st.session_state.image = st.file_uploader(label=" ", accept_multiple_files=False, type=["jpg", "jpeg", "png"])

                if st.session_state.image:
                   col1, col2 = st.columns(2)
                   with col1:
                        st.session_state.complexity = st.selectbox('Complexity Mode Required?*', ['Easy', 'Difficult'], index=0, key="mode1image")
                        st.session_state.no_of_questions = st.number_input('No. of Questions to generate*', key="ai_questions_no_a_image", step=1, max_value=30)
                        st.session_state.mode_of_questions = st.selectbox('Choose Answer Required/Not*', ['Select Option', 'Only Questions', 'Questions with Answers'], index=0, key="quesanswz_image")
                   with col2:
                        st.session_state.type_of_questions = st.selectbox('Choose Question Type*', ['Select Option', 'Short Questions', 'MCQ', 'Fill in the Blanks', 'True and False'], index=0, key="ai_questions_no_p_image")
                        st.session_state.classq = st.selectbox('Choose Class*', ['Select Option', '1', '2', '3', '4', '5', '6', '7', '8', '9', '10', '11', '12'], index=0, key="ai_questions_no_p1_image")
                        st.session_state.cont = st.text_input("Any other context (Optional)", key="eng_image")

                   if st.session_state.mode_of_questions != "Select Option":
                       st.session_state.text_prompt = f'''Based on the image, generate only questions considering following constraints,
                       1. number of questions  - {st.session_state.no_of_questions}
                       2. mode of questions - {st.session_state.mode_of_questions}
                       3. type of questions - {st.session_state.type_of_questions}
                       4. Level of questions - {st.session_state.complexity}
                       5. Class - {st.session_state.classq}
                       6. Question Context - {st.session_state.cont}
                       Generate questions according to Madhya Pradesh School Education Board
                       Response is to be generated in both English and Hindi, first generate in English then in Hindi
                       after generate Answer should be start new line.
                       '''
        
                       image_base64 = to_base64(st.session_state.image)
                       if image_base64:
                          st.write("Image successfully converted to base64")
                          formatted_output = generate_questions(image_base64)
                          img = load_image(st.session_state.image)
                          st.image(img)
                          st.info(formatted_output)
                       else:
                          st.error("Failed to convert image to base64.")
                else:
                    st.info("Please upload an image file.")
   
                
            if choose=="Topic Based Questions":
                
                col1, col2 = st.columns(2)
                with col1:
                    st.session_state.topic_name = st.text_input('Specific Topic Name',placeholder="Topic Name",key="tt")
                    st.session_state.complexity =  st.selectbox('Complexity Mode Required?*', ['Easy', 'Difficult'],index=0,key="mode1")
                    st.session_state.no_of_questions = st.number_input('No. of  Questions to generate*',key="ai_questions_no_a",step=1,max_value=30)
                with col2:
                    st.session_state.type_of_questions =  st.selectbox('Choose Question Type*', ['Select Option','Short Questions','MCQ','Fill in the Blanks','True and False'],index=0,key="ai_questions_no_p")
                    #st.session_state.classq =  st.selectbox('Choose Class*', ['Select Option','1','2','3','4','5','6','7','8','9','10','11','12'],index=0,key="ai_questions_no_p1")
                    st.session_state.mode_of_questions = st.selectbox('Choose Answer Required/Not*', ['Select Option','Only Questions', 'Questions with Answers'],index=0,key="quesanswz")
                if st.button("Submit"):
                   if st.session_state.topic_name and st.session_state.mode_of_questions!='Select Option' :
                      st.session_state.llm = ConversationChain( llm=ChatOpenAI(
                                                              model = "gpt-4o-mini",
                                                              temperature=0.7,
                                                              api_key=openai_api_key2
                                                              ))

                      formatted_output = st.session_state.llm.predict(input = ai_topic_prompt.format(st.session_state.topic_name,
                                                                                                st.session_state.no_of_questions,
                                                                                                    st.session_state.mode_of_questions,
                                                                                                    st.session_state.type_of_questions,
                                                                                                  st.session_state.complexity))
                    
                    
                      st.write(formatted_output)                                                 
                      markdown_to_pdf(formatted_output,'question.pdf')
                    
                    
                      word_doc = create_word_doc(formatted_output)
                      doc_buffer = download_doc(word_doc)
                      st.download_button(label="Download Word Document", 
                      data=doc_buffer, 
                      file_name="generated_document.docx", 
                      mime="application/octet-stream",
                      key='worddownload3')
                    
if st.session_state.teach=='Students':
    choose=st.radio("Select Options",("Pre Uploaded","Ask a Query","Text Analyzer"),horizontal=True)
    if choose == "Pre Uploaded":
        subjects_folder = "./preuploaded"

        # List subjects
        subjects_list = [d for d in os.listdir(subjects_folder) if os.path.isdir(os.path.join(subjects_folder, d))]
        subjects_list.sort()
        subjects_list.insert(0, "Select subject")

        selected_subject = st.selectbox("Select a subject", subjects_list, index=0, key='subject_selector')

        if selected_subject != "Select subject":
            # List chapters for the selected subject
            folder_path = os.path.join(subjects_folder, selected_subject)
            files_list = list_files(folder_path)
            files_list = [remove_extension(filename) for filename in files_list]
            files_list.sort()
            files_list.insert(0, "All Chapters")

            selected_file = st.selectbox("Select a chapter (optional)", files_list, index=0, key='chapter_selector')

            # Initialize text variable
            st.session_state.text = ""

            if selected_file == "All Chapters":
                for file in files_list[1:]:  # Skip "All Chapters" option
                    pdf_file_path = os.path.join(folder_path, file + '.pdf')
                    st.session_state.text += pdf_to_text(pdf_file_path) + "\n"
            elif selected_file != "Select chapter":
                st.session_state.filename = []
                pdf_file_path = os.path.join(folder_path, selected_file + '.pdf')
                st.session_state.filename.append(selected_file)
                st.session_state.text = pdf_to_text(pdf_file_path)

            if st.session_state.text:
                # Input settings
                col1, col2 = st.columns(2)
                with col1:
                    st.session_state.complexity = st.selectbox('Complexity Mode Required?*', ['Easy', 'Difficult'], index=0, key="mode")
                    st.session_state.no_of_questions = st.number_input('No. of Questions to generate*', key="ai_questions", step=1, max_value=30)
                    st.session_state.mode_of_questions = st.selectbox('Choose Answer Required/Not*', ['Only Questions', 'Questions with Answers'], index=0, key="quesansw")
                with col2:
                    st.session_state.type_of_questions = st.selectbox('Choose Question Type*', ['Short Questions', 'Long Questions', 'MCQ', 'Fill in the Blanks', 'True and False'], index=0)
                    st.session_state.language = st.selectbox('Choose Response Language Mode*', ['English', 'English and Hindi'], index=0, key="lang")

                # Button to trigger processing
                if st.button("Submit"):
                    if st.session_state.text and st.session_state.mode_of_questions != 'Select Option':
                        st.session_state.llm = ConversationChain(llm=ChatOpenAI(model="gpt-4o", temperature=0.7, api_key=openai_api_key2))
                        chapter_info = f"Chapter: {selected_file}" if selected_file != "All Chapters" else "All Chapters"
                        formatted_output = st.session_state.llm.predict(input=ai_topic_prompt1.format(
                            chapter_info,
                            st.session_state.no_of_questions,
                            st.session_state.text,
                            st.session_state.language,
                            st.session_state.mode_of_questions,
                            st.session_state.type_of_questions,
                            st.session_state.complexity,
                            st.session_state.no_of_questions
                        ))

                        st.info(formatted_output)
                        markdown_to_pdf(formatted_output, 'question.pdf')
                        word_doc = create_word_doc(formatted_output)
                        doc_buffer = download_doc(word_doc)

                        st.download_button(
                            label="Download Word Document",
                            data=doc_buffer,
                            file_name="generated_document.docx",
                            mime="application/octet-stream",
                            key='worddownload'
                        )            

    if choose=="Text Analyzer":
                
                txt = st.text_area(
                "Text to Generate Questions"
                )
                if len(txt)>0:
                    #st.write(txt)
                    with open('dictionary.json','r') as f:
                            existing_dictionary = json.load(f)

                    lowercase_dict = {key.lower(): value for key, value in existing_dictionary.items()}
                    st.session_state.text=correct_bhashni_translations(txt,lowercase_dict)
                    #st.write(st.session_state.text)

                    col1, col2 = st.columns(2)
                    with col1:
                        #st.session_state.complexity =  st.selectbox('Complexity Mode Required?*', ['No', 'Yes'],index=0,key="mode")
                        st.session_state.no_of_questions = st.number_input('No. of  Questions to generate*',key="ai_questions_no_ks",step=1,max_value=30)
                        st.session_state.mode_of_questions = st.selectbox('Choose Answer Required/Not*', ['Select Option','Only Questions', 'Questions with Answers'],index=0,key="quesansws")
                    with col2:
                        st.session_state.type_of_questions =  st.selectbox('Choose Question Type*', ['Select Option','Short Questions','MCQ','Fill in the Blanks','True and False'],index=0,key="ai_questions_no_ss")
                    if st.button("Submit"):
                        if st.session_state.text and st.session_state.mode_of_questions!='Select Option' :
                            st.session_state.llm = ConversationChain( llm=ChatOpenAI(
                            model_name="gpt-4o-mini",
                            temperature=0.7,
                            api_key=openai_api_key2
                            )) 
                            formatted_output = st.session_state.llm.predict(input = ai_prompt.format(st.session_state.no_of_questions,
                                                                            st.session_state.mode_of_questions,
                                                                            st.session_state.type_of_questions,
                                                                            st.session_state.text))
                            st.info(formatted_output)
                            markdown_to_pdf(formatted_output,'question.pdf')
                            
                            word_doc = create_word_doc(formatted_output)
                            doc_buffer = download_doc(word_doc)
                            st.download_button(label="Download Word Document", 
                                            data=doc_buffer, 
                                            file_name="generated_document.docx", 
                                            mime="application/octet-stream",on_click=None,
                                            key='worddownload2')

    else:
            st.write("")
    if choose=="MCQ Test":
        st.write("In Development Stage")
        st.write('Note: File name should contain subject and class like maths_class10.pdf/.docx')
        files = st.file_uploader('Upload Books,Notes,Question Banks ', accept_multiple_files=True,type=['pdf', 'docx'])
        if files:
            file_extension = files[0].name.split(".")[-1]
            if file_extension == "pdf":
                path = files[0].read()
                name=files[0].name[:-4]
                # Check if the file exists
                if not os.path.exists(name+".txt"):
                    print("File Not Exist")
                    with open(name+'.txt', 'w') as file:
                    # Open a file in write mode (creates a new file if it doesn't exist)
                        st.session_state.text= chat.load_pdf_text(files[0],name)
                        file.write(st.session_state.text)
                        # st.session_state.mcq_chain = ConversationChain( llm=AzureChatOpenAI(
                        # deployment_name="gpt-4turbo",
                        # temperature=0
                        # ))
                        #outputs = chat.mcq_response(st.session_state.text)
                        outputs = generate_quiz(st.session_state.text)
                        #st.info(formatted_output)
                        try:
                            quiz_json = json.loads(outputs[0])["quiz"]
                            st.session_state['quiz_length'] = len(quiz_json)
                            questions = [q['question'] for q in quiz_json]
                            options = [q['options'] for q in quiz_json]
                            answers = [q["answer"] for q in quiz_json]
                            explanations = [q['explanation'] for q in quiz_json]
                            
                            # user selects which question they want to answer
                            # question_num = st.number_input('Choose a question', min_value=1, max_value = len(quiz_json), value=1)
                            question_num = st.session_state['curr_question']
                            answer_choices = options[question_num]
                            col1, col2 = st.columns(2)
                            with col1:
                                st.button('Previous Question', use_container_width=True, on_click=decrement_question_num)
                            with col2:
                                st.button('Next Question', use_container_width=True, on_click=increment_question_num)

                            st.markdown(f"##### Question {question_num + 1} of {len(quiz_json)}: {questions[question_num]}")

                            if 'a' not in st.session_state:
                                st.session_state.a = 0
                                st.session_state.b = 0
                                st.session_state.c = 0
                                st.session_state.d = 0

                            def ChangeA():
                                st.session_state.a,st.session_state.b,st.session_state.c,st.session_state.d = 1,0,0,0
                            def ChangeB():
                                st.session_state.a,st.session_state.b,st.session_state.c,st.session_state.d = 0,1,0,0
                            def ChangeC():
                                st.session_state.a,st.session_state.b,st.session_state.c,st.session_state.d = 0,0,1,0
                            def ChangeD():
                                st.session_state.a,st.session_state.b,st.session_state.c,st.session_state.d = 0,0,0,1
                            def ClearAll():
                                st.session_state.a,st.session_state.b,st.session_state.c,st.session_state.d = 0,0,0,0

                            checkboxA = st.checkbox(answer_choices[0], value = st.session_state.a, on_change = ChangeA)
                            checkboxB = st.checkbox(answer_choices[1], value = st.session_state.b, on_change = ChangeB)
                            checkboxC = st.checkbox(answer_choices[2], value = st.session_state.c, on_change = ChangeC)
                            checkboxD = st.checkbox(answer_choices[3], value = st.session_state.d, on_change = ChangeD)

                            if st.session_state.a:
                                user_answer = answer_choices[0]
                            elif st.session_state.b:
                                user_answer = answer_choices[1]
                            elif st.session_state.c:
                                user_answer = answer_choices[2]
                            elif st.session_state.d:
                                user_answer = answer_choices[3]
                            else:
                                user_answer = None

                            if user_answer is not None:
                                user_answer_num = answer_choices.index(user_answer)
                                if st.button('Submit Answer', type='secondary'):
                                    if user_answer_num == answers[question_num][0]:
                                        st.success(f'Correct! {explanations[question_num]}')
                                    else:
                                        st.error(f'Incorrect :( \n\n The correct answer was: {answer_choices[answers[question_num][0]]}\n\n {explanations[question_num]}')
                        except:
                            st.info('Uh oh... could not generate a quiz for ya! Happy studying!')
                else:
                    print("file Exist")
                    st.session_state.filename=[]
                    with open(name+'.txt', 'r',encoding='ISO-8859-1') as file:
            #        Read the content of the file
                        st.session_state.filename.append(name)
                        st.session_state.text = file.read()
                        #st.write(st.session_state.text)
                        #outputs = generate_quiz(st.session_state.text)
                        # st.session_state.mcq_chain = ConversationChain( llm=AzureChatOpenAI(
                        # deployment_name="gpt-4turbo",
                        # temperature=0
                        # ))
                        #outputs = chat.mcq_response(st.session_state.text)
                        outputs = generate_quiz(st.session_state.text)
                        #st.write(outputs)
                        #print("Output are")
                        #print(outputs)
                        
                        st.write(outputs)
                        corrected_json_data = outputs[0].replace('//', '')
                        quiz_json = json.loads(corrected_json_data)["quiz"]
                        st.session_state['quiz_length'] = len(quiz_json)
                        questions = [q['question'] for q in quiz_json]
                        options = [q['options'] for q in quiz_json]
                        answers = [q["answer"] for q in quiz_json]
                        explanations = [q['explanation'] for q in quiz_json]
                        
                        # user selects which question they want to answer
                        # question_num = st.number_input('Choose a question', min_value=1, max_value = len(quiz_json), value=1)
                        question_num = st.session_state['curr_question']
                        answer_choices = options[question_num]
                        col1, col2 = st.columns(2)
                        with col1:
                            st.button('Previous Question', use_container_width=True, on_click=decrement_question_num)
                        with col2:
                            st.button('Next Question', use_container_width=True, on_click=increment_question_num)

                        st.markdown(f"##### Question {question_num + 1} of {len(quiz_json)}: {questions[question_num]}")

                        if 'a' not in st.session_state:
                            st.session_state.a = 0
                            st.session_state.b = 0
                            st.session_state.c = 0
                            st.session_state.d = 0

                        def ChangeA():
                            st.session_state.a,st.session_state.b,st.session_state.c,st.session_state.d = 1,0,0,0
                        def ChangeB():
                            st.session_state.a,st.session_state.b,st.session_state.c,st.session_state.d = 0,1,0,0
                        def ChangeC():
                            st.session_state.a,st.session_state.b,st.session_state.c,st.session_state.d = 0,0,1,0
                        def ChangeD():
                            st.session_state.a,st.session_state.b,st.session_state.c,st.session_state.d = 0,0,0,1
                        def ClearAll():
                            st.session_state.a,st.session_state.b,st.session_state.c,st.session_state.d = 0,0,0,0

                        checkboxA = st.checkbox(answer_choices[0], value = st.session_state.a, on_change = ChangeA)
                        checkboxB = st.checkbox(answer_choices[1], value = st.session_state.b, on_change = ChangeB)
                        checkboxC = st.checkbox(answer_choices[2], value = st.session_state.c, on_change = ChangeC)
                        checkboxD = st.checkbox(answer_choices[3], value = st.session_state.d, on_change = ChangeD)

                        if st.session_state.a:
                            user_answer = answer_choices[0]
                        elif st.session_state.b:
                            user_answer = answer_choices[1]
                        elif st.session_state.c:
                            user_answer = answer_choices[2]
                        elif st.session_state.d:
                            user_answer = answer_choices[3]
                        else:
                            user_answer = None

                        if user_answer is not None:
                            user_answer_num = answer_choices.index(user_answer)
                            if st.button('Submit Answer', type='secondary'):
                                if user_answer_num == answers[question_num][0]:
                                    st.success(f'Correct! {explanations[question_num]}')
                                else:
                                    st.error(f'Incorrect :( \n\n The correct answer was: {answer_choices[answers[question_num][0]]}\n\n {explanations[question_num]}')


# Function to split text into smaller chunks


    openai_api_key = os.getenv("OPENAI_API_KEY")
    openai_api_key2 = st.secrets["secret_section"]["OPENAI_API_KEY"]

    if 'history' not in st.session_state:
        st.session_state.history = []

    col_1, col_2 = st.columns([2, 1])

    with col_1:
        if choose == "Ask a Query":
            uploaded_file = st.file_uploader("Upload a PDF file", type=["pdf"])

            if uploaded_file is not None:
                pdfreader = PdfReader(uploaded_file)
                raw_text = ''
                for i, page in enumerate(pdfreader.pages):
                    content = page.extract_text()
                    if content:
                        raw_text += content

                text_splitter = CharacterTextSplitter(
                    separator="\n",
                    chunk_size=800,
                    chunk_overlap=200,
                    length_function=len,
                )
                texts = text_splitter.split_text(raw_text)
                #st.write(f"PDF loaded and split into {len(texts)} chunks.")

                embeddings = OpenAIEmbeddings(api_key=openai_api_key2)
                document_search = FAISS.from_texts(texts, embeddings)
                #st.write("Document embeddings created and stored in FAISS index.")

                chain = load_qa_chain(OpenAI(api_key=openai_api_key2), chain_type="stuff")

                query = st.chat_input("Ask a question about the PDF:")
                st.write(query)
                if query:
                    docs = document_search.similarity_search(query)
                    answer = chain.run(input_documents=docs, question=query)
                    st.session_state.history.append((query, answer))
                    st.write("Answer:", answer)

                    # Translate to Hindi
                    translator = Translator(to_lang="hi")
                    
                    # Split query and answer into smaller chunks for translation
                    query_chunks = split_text_into_chunks(query, 500)
                    query_hindi_chunks = [translator.translate(chunk) for chunk in query_chunks]
                    query_hindi = " ".join(query_hindi_chunks)

                    answer_chunks = split_text_into_chunks(answer, 500)
                    answer_hindi_chunks = [translator.translate(chunk) for chunk in answer_chunks]
                    answer_hindi = " ".join(answer_hindi_chunks)

                    st.session_state.history[-1] += (query_hindi, answer_hindi)
                    st.write("**In Hindi:**")
                    st.write(f"**Q:** {query_hindi}")
                    st.write(f"**A:** {answer_hindi}")

                if st.session_state.history:
                    doc = Document()
                    doc.add_heading('Questions and Answers', 0)

                    for i, (question, answer, question_hindi, answer_hindi) in enumerate(st.session_state.history):
                        doc.add_heading(f"Q{i+1}: {question}", level=1)
                        doc.add_paragraph(f"A{i+1}: {answer}")
                        doc.add_heading(f"Q{i+1} (Hindi): {question_hindi}", level=1)
                        doc.add_paragraph(f"A{i+1} (Hindi): {answer_hindi}")

                    # Save the document in the current directory with a unique name
                    timestamp = time.strftime("%Y%m%d-%H%M%S")
                    doc_path = f"QnA_History_{timestamp}.docx"
                    doc.save(doc_path)

                    with open(doc_path, "rb") as f:
                        st.download_button(
                            label="Download Word Document",
                            data=f,
                            file_name=doc_path,
                            mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document"
                        )

            # Add a section for predefined prompts
            

            with col_2:
                st.write("### History")
                for i, (question, answer, question_hindi, answer_hindi) in enumerate(st.session_state.history):
                    st.write(f"**Q{i+1}:** {question}")
                    st.write(f"**A{i+1}:** {answer}")
                    st.write(f"**Q{i+1} (Hindi):** {question_hindi}")
                    st.write(f"**A{i+1} (Hindi):** {answer_hindi}")

# Add a download button to download the history as a Word document
    
    if choose=="Terminologies and Keyterms":
        st.write('Note: File name should contain subject and class like maths_class10.pdf/.docx')
        files = st.file_uploader('Upload Books,Notes,Question Banks ', accept_multiple_files=True,type=['pdf', 'docx'])
        if files:
            file_extension = files[0].name.split(".")[-1]
            if file_extension == "pdf":
                path = files[0].read()
                name=files[0].name[:-4]
                #Check if the file exists
                if not os.path.exists(name+".txt"):
                    print("File Not Exist")
                    with open(name+'.txt', 'w') as file:
                    #Open a file in write mode (creates a new file if it doesn't exist)
                        st.session_state.text= chat.load_pdf_text(files[0],name)
                        file.write(st.session_state.text)
                else:
                    print("file Exist")
                    st.session_state.filename=[]
                    with open(name+'.txt', 'r',encoding='ISO-8859-1') as file:
                    #Read the content of the file
                        st.session_state.filename.append(name)
                        st.session_state.text = file.read()
                        st.session_state.mcq_chain = ConversationChain( llm=ChatOpenAI(
                        model="gpt-3.5-turbo",
                        temperature=0.7
                        ))
                        outputs = chat.mcq_response(st.session_state.text)
                        st.write(outputs)
                        markdown_to_pdf(outputs,'question.pdf')
                        
                        word_doc = create_word_doc(outputs)
                        doc_buffer = download_doc(word_doc)
                        st.download_button(label="Download Word Document", 
                                        data=doc_buffer, 
                                        file_name="generated_document.docx", 
                                        mime="application/octet-stream",
                                        key='worddownload3')
if st.session_state.teach == 'Administration':
    if 'selected_file' not in st.session_state:
            st.session_state.selected_file = None

    # Define the options and handle the session state for the selected option
    if 'selected_option' not in st.session_state:
        st.session_state.selected_option = None

    def reset_selected_file():
        st.session_state.selected_file = None

    st.session_state.selected_option = st.radio("Select Options", ("Add Document", "Download Document", "Delete Document", "View Documents"), horizontal=True, on_change=reset_selected_file)

    choose = st.session_state.selected_option

    base_folder = "./preuploaded"

    if choose == "Add Document":
        subject_folders = [d for d in os.listdir(base_folder) if os.path.isdir(os.path.join(base_folder, d))]
        subject_folders.sort()
        subject_folders.insert(0, "Select Subject")
        
        selected_subject = st.selectbox("Select Subject", subject_folders, index=0, key='subject_selector')

        if selected_subject != "Select Subject":
            files = st.file_uploader('Upload Books, Notes, Question Banks', accept_multiple_files=True, type=['pdf'])
            if files:
                subject_folder_path = os.path.join(base_folder, selected_subject)
                os.makedirs(subject_folder_path, exist_ok=True)
                
                for uploaded_file in files:
                    file_path = os.path.join(subject_folder_path, uploaded_file.name)
                    if not os.path.exists(file_path):
                        with open(file_path, 'wb') as file:
                            file.write(uploaded_file.getbuffer())
                        st.success(f"{uploaded_file.name} uploaded successfully.")
                    else:
                        st.warning(f"{uploaded_file.name} already exists.")

    elif choose == "Download Document":
        subject_folders = [d for d in os.listdir(base_folder) if os.path.isdir(os.path.join(base_folder, d))]
        subject_folders.sort()
        subject_folders.insert(0, "Select Subject")
        
        selected_subject = st.selectbox("Select Subject", subject_folders, index=0, key='download_subject_selector')

        if selected_subject != "Select Subject":
            subject_folder_path = os.path.join(base_folder, selected_subject)
            pdf_files = [file for file in os.listdir(subject_folder_path) if file.endswith(".pdf")]
            pdf_files.insert(0, "Select Document")
            selected_file = st.selectbox("Select Document", pdf_files, index=0, key='download_selected_file')

            if selected_file != "Select Document":
                file_path = os.path.join(subject_folder_path, selected_file)
                with open(file_path, "rb") as file:
                    st.download_button(label="Download", data=file, file_name=selected_file, mime="application/pdf")
            else:
                st.info("Select a document to download.")

    elif choose == "View Documents":
        subject_folders = [d for d in os.listdir(base_folder) if os.path.isdir(os.path.join(base_folder, d))]
        subject_folders.sort()
        subject_folders.insert(0, "Select Subject")

        selected_subject = st.selectbox("Select Subject", subject_folders, index=0, key='view_subject_selector')

        if selected_subject != "Select Subject":
            subject_folder_path = os.path.join(base_folder, selected_subject)
            pdf_files = [file for file in os.listdir(subject_folder_path) if file.endswith(".pdf")]
            if pdf_files:
                st.write(f"Documents in the folder for {selected_subject}:")
                for pdf_file in pdf_files:
                    st.write(pdf_file)
            else:
                st.info("No documents found in the selected subject folder.")

    elif choose == "Delete Document":
        subject_folders = [d for d in os.listdir(base_folder) if os.path.isdir(os.path.join(base_folder, d))]
        subject_folders.sort()
        subject_folders.insert(0, "Select Subject")

        selected_subject = st.selectbox("Select Subject", subject_folders, index=0, key='delete_subject_selector')

        if selected_subject != "Select Subject":
            subject_folder_path = os.path.join(base_folder, selected_subject)
            pdf_files = [file for file in os.listdir(subject_folder_path) if file.endswith(".pdf")]
            pdf_files.insert(0, "Select Document")
            selected_file = st.selectbox("Select Document", pdf_files, index=0, key='delete_selected_file')

            if selected_file != "Select Document":
                os.remove(os.path.join(subject_folder_path, selected_file))
                st.success(f"{selected_file} has been successfully removed.")
            else:
                st.info("Select a document to delete.")
    if choose=="Add Word to Dictionary":
        if st.button("View Dictionary"):
            with open('dictionary.json','r') as f:
                    existing_dictionary = json.load(f)
            lowercase_dict = {key.lower(): value for key, value in existing_dictionary.items() if isinstance(value, str)}
            st.write(lowercase_dict)

        # Input field for searching a word
        search_word = st.text_input('Enter English Word to Search:', '')
        if search_word:
            with open('dictionary.json','r') as f:
                existing_dictionary = json.load(f)
            translation = existing_dictionary.get(search_word, 'Word not found')
            st.write(f"Hindi Translation for '{search_word}': {translation}")

        with open('dictionary.json','r') as f:
            st.session_state.existing_dictionary = json.load(f)

        # Function to save dictionary to a local file
        def save_dictionary_to_file(filename):
            #st.write(st.session_state.existing_dictionary)
            with open(filename, 'w') as f:
                json.dump(st.session_state.existing_dictionary, f)

        col1, col2 = st.columns(2)
        with col1:

            st.write("##### English Words (Separated by commas)")
            eng_txt = st.text_input(
                "Enter English Words",
                key="eng"
            )
        with col2:

            st.write("##### Hindi Words (Separated by commas)")
            punjabi_txt = st.text_input(
                "Enter Hindi Words",
                key="hi"
            )

        lowercase_dict = {key.lower(): value for key, value in st.session_state.existing_dictionary.items()}

        if st.button("Request to Add"):
        # Split the entered English and Punjabi words by commas and remove leading/trailing whitespace
            eng_words = [word.strip() for word in eng_txt.split(",")]
            punjabi_words = [word.strip() for word in punjabi_txt.split(",")]

            # Add each pair of words to the lowercase dictionary
            for eng_word, punjabi_word in zip(eng_words, punjabi_words):
                # Convert entered English word to lowercase
                eng_lower = eng_word.lower()
                # Add the entry to the lowercase dictionary
                st.session_state.existing_dictionary[eng_lower] = punjabi_word
            st.success("Words added to dictionary!")
            st.write("Dictionary:", st.session_state.existing_dictionary)
            save_dictionary_to_file("dictionary.json")
            st.success("Dictionary saved to local!")
                    
footer = """
    <style>
    body {
        margin: 0;
        padding-top: 70px;  /* Add padding to prevent content from being hidden behind the footer */
    }
    .footer {
        position: absolute;
        top: 80px;
        left: 0;
        width: 100%;
        background-color: #002F74;
        color: white;
        text-align: center;
        padding: 5px;
        font-weight: bold;
        z-index: 1000;  /* Ensure it is on top of other elements */
        display: flex;
        align-items: center;
        justify-content: space-between;
        flex-wrap: wrap;
    }
    .footer p {
        font-style: italic;
        font-size: 14px;
        margin: 0;
        flex: 1 1 50%;  /* Flex-grow, flex-shrink, flex-basis */
    }
    @media (max-width: 600px) {
        .footer p {
            flex-basis: 100%;
            text-align: center;
            padding-top: 10px;
        }
    }
    </style>
    <div class="footer">
        <p style="text-align: left;">Copyright © 2024 MPSeDC. All rights reserved.</p>
        <p style="text-align: right;">The responses provided on this website are AI-generated. User discretion is advised.</p>
    </div>
"""

st.markdown(footer, unsafe_allow_html=True)
